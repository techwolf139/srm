from docx import Document
from typing import List, Dict, Tuple
import re

RISK_KEYWORDS = {
    'CRITICAL': [
        # 无限责任/无上限
        '无限责任', 'unlimited liability', '全权负责', '无限制', '无上限', '无限额',
        # 自动续约
        '自动续约', 'auto-renewal', '自动延续', '期满自动',
        # 单方权利
        '单方终止', '单方面解除', '单方解除', '任意终止', '随时终止',
        # 转让限制
        '未经同意', 'without consent', '未经书面同意', '不得转让',
        # 诉讼权利放弃
        '放弃陪审', '放弃诉讼', '放弃仲裁', '放弃抗辩',
        'class action', '集体诉讼', '集团诉讼',
        # 永久有效
        '永久有效', '永久保密', '永久有效期', '永久约束',
        # 不可撤销
        '不可撤销', '不可撤回', '无条件执行',
        # 独家/排他
        '独家合作', '排他性', '独家供应商', '唯一供应商',
        # 强制执行
        '强制执行', '强制执行权', '强制赔偿',
    ],
    'HIGH': [
        # 模糊努力标准
        '最大努力', 'best efforts', 'reasonable efforts', '合理努力', '尽一切努力',
        '一切合理努力', '商业上合理',
        # 模糊定义
        'material', '重大', '及时', '合理时间', '适当',
        '主要', '关键', '重要', '实质性',
        # 赔偿/ Indemnification
        '赔偿', 'indemnify', 'indemnification', '全额赔偿', '全部赔偿',
        '承担一切', '承担所有',
        # 竞业禁止
        '竞业', 'non-compete', '竞争禁止', '禁止竞争',
        # 知识产权
        '知识产权', 'intellectual property', 'IP', '版权', '专利权',
        '著作权', '技术成果', '职务成果',
        # 最惠国
        '最惠国', 'most favored', 'MFN', '最惠待遇',
        # 价格调整
        '价格调整', 'price increase', '价格变动', '费率调整',
        '涨价', '加价',
        # 最低消费
        '最低消费', '最低费用', 'minimum', '不少于',
        # 违约金
        '违约金', 'penalty', '罚金', '违约罚则',
        # 排他性
        '排他性', '独家', '独占',
        # 优先购买权
        '优先购买权', '优先权', 'first right',
        # 保密义务
        '保密义务', 'confidentiality', '保密信息', '商业秘密',
        # 转让权
        '转让权', 'assignment', '权利转让', '义务转让',
        # 担保
        '担保', 'guarantee', '保证责任', '连带保证',
    ],
    'MEDIUM': [
        # 通知期限
        '通知期限', '提前通知', '书面通知',
        # 验收标准
        '验收标准', '验收程序', '验收合格',
        # 付款条件
        '付款条件', '支付方式', '结算方式',
        # 质量保证
        '质量保证', '质量标准', '保修',
        # 不可抗力
        '不可抗力', 'force majeure', '不可预见',
        # 争议解决
        '争议解决', 'dispute resolution', '仲裁条款', '管辖权',
        # 适用法律
        '适用法律', 'governing law', '准据法',
        # 终止条件
        '终止条件', '解除条件', '终止条款',
        # 变更条款
        '变更条款', '修改条款', '变更权',
        # 责任限制
        '责任限制', 'limitation of liability', '责任上限',
        # 损失赔偿
        '损失赔偿', '损害赔偿', '赔偿责任',
    ]
}

FINANCIAL_RISK_PATTERNS = {
    'EXCESSIVE_PENALTY': r'违约金[为是]\s*([\d.]+)\s*%',
    'NO_PENALTY_CAP': r'违约金.*无上限|无限制.*违约金',
    'HIGH_INTEREST': r'利率\s*([\d.]+)\s*%',
    'NO_LIABILITY_CAP': r'责任.*无上限|无限.*责任',
    'PERCENTAGE_CLAUSE': r'([\d.]+)\s*%\s*日|每日\s*([\d.]+)\s*%',
}

DATA_PROTECTION_KEYWORDS = [
    '个人信息', 'personal information', '数据保护', 'data protection',
    '隐私政策', 'privacy', '个人信息保护', '数据安全',
    '跨境传输', 'cross-border', '数据出境',
    '敏感信息', 'sensitive information', '个人隐私',
]

LABOR_LAW_KEYWORDS = [
    '试用期', '试用期过长', '试用期工资',
    '社会保险', '社保', '五险一金',
    '加班费', '加班工资', '工时',
    '年假', '带薪年假', '年休假',
    '工资支付', '发薪日', '薪酬',
    '解除劳动合同', '经济补偿', '赔偿金',
    '竞业限制', '服务期', '培训费',
]

def extract_amounts(text: str) -> List[Dict]:
    results = []
    amount_patterns = [
        (r'¥\s*([\d,]+(?:\.\d{2})?)', 'CNY'),
        (r'RMB\s*([\d,]+(?:\.\d{2})?)', 'CNY'),
        (r'人民币\s*([\d,]+(?:\.\d{2})?)', 'CNY'),
        (r'\$([\d,]+(?:\.\d{2})?)', 'USD'),
        (r'USD\s*([\d,]+(?:\.\d{2})?)', 'USD'),
        (r'([\d,]+(?:\.\d{2})?)\s*万元', 'WAN_CNY'),
        (r'([\d,]+(?:\.\d{2})?)\s*元', 'CNY'),
        (r'([\d,]+(?:\.\d{2})?)\s*美元', 'USD'),
    ]
    for pattern, currency in amount_patterns:
        for match in re.finditer(pattern, text):
            value_str = match.group(1).replace(',', '')
            try:
                value = float(value_str)
                if currency == 'WAN_CNY':
                    value = value * 10000
                    currency = 'CNY'
                results.append({
                    'value': value,
                    'currency': currency,
                    'raw': match.group(0),
                    'context': text[max(0, match.start()-20):match.end()+20]
                })
            except ValueError:
                pass
    return results

def extract_percentages(text: str) -> List[Dict]:
    results = []
    percentage_patterns = [
        r'([\d,]+(?:\.\d+)?)\s*%',
        r'百分之([\d,]+(?:\.\d+)?)',
    ]
    for pattern in percentage_patterns:
        for match in re.finditer(pattern, text):
            value_str = match.group(1).replace(',', '')
            try:
                value = float(value_str)
                results.append({
                    'value': value,
                    'raw': match.group(0),
                    'context': text[max(0, match.start()-30):match.end()+30]
                })
            except ValueError:
                pass
    return results

def detect_financial_risks(doc: Document) -> List[Dict]:
    all_text = extract_all_text(doc)
    risks = []
    
    for risk_name, pattern in FINANCIAL_RISK_PATTERNS.items():
        matches = re.finditer(pattern, all_text)
        for match in matches:
            if risk_name == 'EXCESSIVE_PENALTY':
                pct = float(match.group(1))
                if pct > 30:
                    risks.append({
                        'type': risk_name,
                        'level': 'HIGH',
                        'value': pct,
                        'context': all_text[max(0, match.start()-50):match.end()+50],
                        'warning': f'违约金比例{pct}%可能过高（超过30%可能被法院调减）'
                    })
            elif risk_name == 'HIGH_INTEREST':
                pct = float(match.group(1))
                if pct > 15.4:
                    risks.append({
                        'type': risk_name,
                        'level': 'CRITICAL',
                        'value': pct,
                        'context': all_text[max(0, match.start()-50):match.end()+50],
                        'warning': f'利率{pct}%超过LPR四倍上限（可能违法）'
                    })
    
    return risks

def detect_data_protection_risks(doc: Document) -> List[Dict]:
    all_text = extract_all_text(doc)
    risks = []
    found_positions = set()
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text_lower = para.text.lower()
        for keyword in DATA_PROTECTION_KEYWORDS:
            if keyword.lower() in para_text_lower:
                pos_key = (para_idx, keyword)
                if pos_key not in found_positions:
                    found_positions.add(pos_key)
                    risks.append({
                        'keyword': keyword,
                        'location': f'第{para_idx+1}段',
                        'context': para.text[:150],
                        'concern': '涉及个人信息/数据处理，需确认是否符合《个人信息保护法》要求'
                    })
    
    return risks

def detect_labor_law_risks(doc: Document) -> List[Dict]:
    all_text = extract_all_text(doc)
    risks = []
    found_positions = set()
    
    for para_idx, para in enumerate(doc.paragraphs):
        para_text_lower = para.text.lower()
        for keyword in LABOR_LAW_KEYWORDS:
            if keyword.lower() in para_text_lower:
                pos_key = (para_idx, keyword)
                if pos_key not in found_positions:
                    found_positions.add(pos_key)
                    risks.append({
                        'keyword': keyword,
                        'location': f'第{para_idx+1}段',
                        'context': para.text[:150],
                        'concern': '涉及劳动法条款，需确认是否符合《劳动合同法》规定'
                    })
    
    return risks

def analyze_financial_terms(doc: Document) -> Dict:
    all_text = extract_all_text(doc)
    all_amounts = extract_amounts(all_text)
    all_percentages = extract_percentages(all_text)
    
    total_amount = 0
    currencies = {}
    for amt in all_amounts:
        total_amount += amt['value']
        curr = amt['currency']
        currencies[curr] = currencies.get(curr, 0) + amt['value']
    
    high_percentages = [p for p in all_percentages if p['value'] > 30]
    
    return {
        'total_mentioned': total_amount,
        'currencies': currencies,
        'amount_count': len(all_amounts),
        'percentage_count': len(all_percentages),
        'percentages': all_percentages[:20],
        'high_percentages': high_percentages,
        'high_amounts': [a for a in all_amounts if a['value'] > 1000000][:10],
        'flagged_unlimited': '无限' in all_text or '无限制' in all_text or '无上限' in all_text,
        'financial_risks': detect_financial_risks(doc),
        'data_protection_risks': detect_data_protection_risks(doc),
        'labor_law_risks': detect_labor_law_risks(doc),
    }

def create_financial_summary(doc: Document) -> str:
    analysis = analyze_financial_terms(doc)
    lines = ['## 合同财务分析\n']
    lines.append(f'**金额提及次数:** {analysis["amount_count"]}')
    lines.append(f'**百分比提及次数:** {analysis["percentage_count"]}')
    lines.append(f'**涉及总金额:** {analysis["total_mentioned"]:,.2f} 元')
    
    if analysis['currencies']:
        lines.append('\n**币种分布:**')
        for curr, amt in analysis['currencies'].items():
            lines.append(f'- {curr}: {amt:,.2f}')
    
    if analysis['high_percentages']:
        lines.append('\n**⚠️ 高风险百分比(>30%):**')
        for p in analysis['high_percentages'][:5]:
            lines.append(f'- {p["raw"]} (可能过高)')
    
    if analysis['high_amounts']:
        lines.append('\n**大额条款(>100万):**')
        for a in analysis['high_amounts'][:5]:
            lines.append(f'- {a["raw"]}: {a["value"]:,.2f} 元')
    
    if analysis['flagged_unlimited']:
        lines.append('\n⚠️ **风险提示: 合同中存在"无限"相关表述，需核实是否有责任上限条款**')
    
    if analysis['financial_risks']:
        lines.append('\n**财务风险检测:**')
        for fr in analysis['financial_risks'][:5]:
            lines.append(f'- {fr["warning"]}')
    
    if analysis['data_protection_risks']:
        lines.append(f'\n**数据保护风险:** {len(analysis["data_protection_risks"])}处')
        for dp in analysis['data_protection_risks'][:3]:
            lines.append(f'- {dp["keyword"]}: {dp["concern"]}')
    
    if analysis['labor_law_risks']:
        lines.append(f'\n**劳动法风险:** {len(analysis["labor_law_risks"])}处')
        for ll in analysis['labor_law_risks'][:3]:
            lines.append(f'- {ll["keyword"]}: {ll["concern"]}')
    
    return '\n'.join(lines)

def read_docx(filepath: str) -> Document:
    return Document(filepath)

def extract_all_text(doc: Document) -> str:
    return '\n'.join(p.text for p in doc.paragraphs if p.text.strip())

def extract_tables(doc: Document) -> List[List[List[str]]]:
    return [[[c.text.strip() for c in r.cells] for r in t.rows] for t in doc.tables]

def detect_risk_keywords(doc: Document) -> Dict[str, List[Dict]]:
    results = {level: [] for level in RISK_KEYWORDS}
    for para_idx, para in enumerate(doc.paragraphs):
        text = para.text.lower()
        for level, keywords in RISK_KEYWORDS.items():
            for keyword in keywords:
                if keyword.lower() in text:
                    results[level].append({
                        'keyword': keyword,
                        'location': f'第{para_idx + 1}段',
                        'context': para.text[:200]
                    })
    return results

def create_audit_report(title: str, findings: List[Dict], output_path: str) -> None:
    doc = Document()
    doc.add_heading(title, level=1)
    
    critical = sum(1 for f in findings if f['level'] == 'CRITICAL')
    high = sum(1 for f in findings if f['level'] == 'HIGH')
    medium = sum(1 for f in findings if f['level'] == 'MEDIUM')
    
    doc.add_paragraph(f'审计汇总:\n- 严重问题: {critical}\n- 高风险问题: {high}\n- 中等问题: {medium}')
    doc.add_page_break()
    
    for level_name, level_key in [('严重问题', 'CRITICAL'), ('高风险问题', 'HIGH'), ('中等问题', 'MEDIUM')]:
        items = [f for f in findings if f['level'] == level_key]
        if items:
            doc.add_heading(level_name, level=2)
            table = doc.add_table(rows=len(items) + 1, cols=4)
            table.style = 'Table Grid'
            for i, h in enumerate(['位置', '条款摘要', '风险说明', '建议']):
                cell = table.rows[0].cells[i]
                cell.text = h
                if cell.paragraphs[0].runs:
                    cell.paragraphs[0].runs[0].bold = True
            for row_idx, item in enumerate(items, 1):
                table.rows[row_idx].cells[0].text = item.get('location', '')
                table.rows[row_idx].cells[1].text = item.get('clause', '')
                table.rows[row_idx].cells[2].text = item.get('risk', '')
                table.rows[row_idx].cells[3].text = item.get('recommendation', '')
            doc.add_paragraph()
    
    doc.add_heading('审计结论', level=2)
    if critical > 0:
        doc.add_paragraph('建议重新谈判或拒绝签署')
    elif high > 0:
        doc.add_paragraph('建议与对方协商修改相关条款')
    else:
        doc.add_paragraph('合同条款可接受')
    
    doc.save(output_path)

def comprehensive_audit(doc: Document) -> Dict:
    risks = detect_risk_keywords(doc)
    finances = analyze_financial_terms(doc)
    
    findings = []
    
    for level in ['CRITICAL', 'HIGH', 'MEDIUM']:
        for r in risks[level][:20]:
            findings.append({
                'level': level,
                'location': r['location'],
                'clause': r['context'][:100],
                'risk': f'{level} risk: {r["keyword"]}',
                'recommendation': '立即修改' if level == 'CRITICAL' else '建议协商'
            })
    
    for fr in finances.get('financial_risks', []):
        findings.append({
            'level': fr['level'],
            'location': '财务条款',
            'clause': fr.get('context', '')[:100],
            'risk': fr['warning'],
            'recommendation': '需修改'
        })
    
    for dp in finances.get('data_protection_risks', []):
        findings.append({
            'level': 'HIGH',
            'location': dp['location'],
            'clause': dp['context'][:100],
            'risk': dp['concern'],
            'recommendation': '需确认合规'
        })
    
    for ll in finances.get('labor_law_risks', []):
        findings.append({
            'level': 'MEDIUM',
            'location': ll['location'],
            'clause': ll['context'][:100],
            'risk': ll['concern'],
            'recommendation': '需确认合法性'
        })
    
    return {
        'risks': risks,
        'finances': finances,
        'findings': findings,
        'summary': create_financial_summary(doc)
    }

if __name__ == '__main__':
    print('合同审计Docx辅助工具 - 增强版')
    print('Functions: read_docx, detect_risk_keywords, analyze_financial_terms, comprehensive_audit')
